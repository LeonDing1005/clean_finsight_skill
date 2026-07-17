from typing import List, Dict, Any, Tuple
import asyncio
import os
import re
import copy
import subprocess
import numpy as np
from datetime import datetime
from pathlib import Path
try:
    import docx2pdf
    _HAS_DOCX2PDF = True
except ImportError:
    _HAS_DOCX2PDF = False
from src.agents.base_agent import BaseAgent
from src.agents import DeepSearchAgent
from src.tools.web.web_crawler import ClickResult
from src.tools import ToolResult, get_tool_categories, get_tool_by_name
from src.agents.report_generator.report_class import Report, Section
from src.utils.helper import extract_markdown, get_md_img
from src.utils.index_builder import IndexBuilder
from src.utils.figure_helper import draw_kline_chart
from src.typography import LATIN_FONT, normalize_docx_typography

def _inject_word_toc(docx_path: str):
    """Inject a real Word TOC field (with updatable page numbers).

    Pandoc --toc creates only a static list; this inserts a genuine Word
    TOC field that users can update via F9 in Word.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(docx_path)
    if not doc.paragraphs:
        return

    first_para = doc.paragraphs[0]._element

    # --- insert a TOC heading paragraph before the first paragraph ---
    def _add_paragraph_before(target_elem, text, bold=False):
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        if bold:
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        target_elem.addprevious(p)
        return p

    _add_paragraph_before(first_para, '目录', bold=True)  # 目录

    # --- inject the real TOC field paragraph ---
    toc_p = OxmlElement('w:p')

    def _make_fld_char(ftype):
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), ftype)
        r = OxmlElement('w:r')
        r.append(fc)
        return r

    toc_p.append(_make_fld_char('begin'))
    instr_run = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run.append(instr)
    toc_p.append(instr_run)
    toc_p.append(_make_fld_char('separate'))
    toc_p.append(_make_fld_char('end'))

    first_para.addprevious(toc_p)

    # --- left-align the Reference Data Sources section ---
    in_ref = False
    for para in doc.paragraphs:
        if para.text.strip() == 'Reference Data Sources' and para.style.name.startswith('Heading'):
            in_ref = True
            continue
        if in_ref:
            if para.style.name.startswith('Heading'):
                break
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(docx_path)


class ReportGenerator(BaseAgent):
    AGENT_NAME = 'report_generator'
    AGENT_DESCRIPTION = 'a agent that can generate report from the data'
    NECESSARY_KEYS = ['task']
    def __init__(
        self,
        config,
        tools = [],
        use_llm_name: str = "deepseek-chat",
        use_embedding_name: str = "qwen3-embedding",
        enable_code = False,
        memory = None,
        agent_id: str = None
    ):
        super().__init__(
            config=config,
            tools=tools,
            use_llm_name=use_llm_name,
            enable_code=enable_code,
            memory=memory,
            agent_id=agent_id
        )
        # Load prompts based on language and target type settings
        from src.utils.prompt_loader import get_prompt_loader
        
        target_language = self.config.config.get('language', 'zh')
        language_mapping = {
            'zh': 'Chinese (中文)',
            'en': 'English'
        }
        target_language_name = language_mapping.get(target_language, target_language)
        self.target_language_name = target_language_name
        target_type = self.config.config.get('target_type', 'general')
        
        
        # Load prompts using the new YAML-based loader
        self.prompt_loader = get_prompt_loader('report_generator', report_type=target_type)
        
        # Store prompts as instance attributes for easy access
        self.SECTION_WRITING_PROMPT = self.prompt_loader.get_prompt('section_writing')
        self.SECTION_WRITING_WO_CHART_PROMPT = self.prompt_loader.get_prompt('section_writing_wo_chart')
        self.FINAL_POLISH_PROMPT = self.prompt_loader.get_prompt('final_polish')
        
        # For general reports, use outline_draft; for financial, use outline_draft as well
        # (both YAML files have 'outline_draft' key)
        self.DRAFT_GENERATOR_PROMPT = self.prompt_loader.get_prompt('outline_draft')
        
        self.CRITIQUE_PROMPT = self.prompt_loader.get_prompt('outline_critique')
        self.REFINEMENT_PROMPT = self.prompt_loader.get_prompt('outline_refinement')
        
        # used for adding abstract and title
        self.TITLE_PROMPT = self.prompt_loader.get_prompt('title_generation')
        self.ABSTRACT_PROMPT = self.prompt_loader.get_prompt('abstract')

        # used for cover page
        self.TABLE_BEAUTIFY_PROMPT = self.prompt_loader.get_prompt('table_beautify')
        
        self.use_embedding_name = use_embedding_name
        # Phase checkpoints: outline → sections → post_process
        self._phase: str = 'outline'
        # Section-level progress counter
        self._section_index_done: int = 0
        # Post-process sub-stages: 0-image, 1-abstract/title, 2-cover, 3-reference, 4-render
        self._post_stage: int = 0
        

    def _set_default_tools(self):
        """
        Attach the default tools/agents required by the report generator.
        """
        tool_list = []
        # Attach the deep-search agent (sharing the same memory)
        tool_list.append(DeepSearchAgent(config=self.config, use_llm_name=self.use_llm_name, memory=self.memory))
        for tool in tool_list:
            self.memory.add_dependency(tool.id, self.id)
        self.tools = tool_list
    
    async def _prepare_executor(self):
        """
        Prepare the code executor with data access functions for section writing.
        """
        if not self.enable_code:
            return
        current_task_data = self.current_task_data
        tool_list = self.tools
        collect_data_list = self.memory.get_collect_data(exclude_type=['search', 'click'])
        analysis_result_list = self.memory.get_analysis_result()
        
        def _get_data(data_id: int):
            """Get dataset by index"""
            if 0 <= data_id < len(collect_data_list):
                return collect_data_list[data_id].data
            else:
                raise ValueError(f"Invalid data_id: {data_id}. Available range: 0-{len(collect_data_list)-1}")
        
        def _get_analysis_result(data_id: int):
            """Get analysis results matching the query"""
            # Use LLM-based selection to find relevant analysis results
            if 0 <= data_id < len(analysis_result_list):
                return str(analysis_result_list[data_id])[:3000]
            else:
                raise ValueError(f"Invalid data_id: {data_id}. Available range: 0-{len(analysis_result_list)-1}")
        
        def _get_deepsearch_result(query: str):
            """Call deep search agent.

            Uses the async bridge to avoid the asyncio.run() deadlock when
            called from inside exec()'d code within a running event loop.
            """
            from src.utils.async_bridge import get_async_bridge
            bridge = get_async_bridge()
            ds_agent = tool_list[0]
            output = bridge.run_async(ds_agent.async_run(input_data={
                'task': current_task_data.get('task', ''),
                'query': query
            }))
            return output['final_result']
        
        self.code_executor.set_variable("get_data", _get_data)
        self.code_executor.set_variable("get_analysis_result", _get_analysis_result)
        self.code_executor.set_variable("get_data_from_deep_search", _get_deepsearch_result)
        
    
    async def _prepare_init_prompt(self, input_data: dict) -> list[dict]:
        task = input_data.get('task')
        section_outline = input_data.get('section_outline')
        max_iterations = input_data.get('max_iterations', 10)
        if not task:
            raise ValueError("Input data must contain a 'task' key.")
        
        # Get data API description from prompts
        data_api_description = self.prompt_loader.get_prompt('data_api')
        
        # Prepare data information for the agent
        collect_data_list = self.memory.get_collect_data(exclude_type=['search', 'click'])
        analysis_result_list = self.memory.get_analysis_result()
        data_info = "\n\n## Available Datas\n\n"
        for idx, item in enumerate(collect_data_list):
            data_info += f"**Data ID {idx}:**\n{item.brief_str()}\n\n"
        data_info += "\nYou can access these datasets using `get_data(data_id)` in your code.\n"
        data_info += "\n\n## Available Analysis Reports\n\n"
        for idx, item in enumerate(analysis_result_list):
            data_info += f"**Analysis Report ID {idx}:**\n{item.brief_str()}\n\n"
        data_info += "\nYou can access these analysis reports using `get_analysis_result(analysis_result_id)` in your code.\n"
        
        if self.enable_chart:
            # Build reference strings for prompt placeholders
            ref_data_parts = []
            for idx, item in enumerate(collect_data_list):
                ref_data_parts.append(f"Data ID {idx}: {item.brief_str()}")
            reference_data = "\n".join(ref_data_parts) if ref_data_parts else "No collected data available"

            ref_analysis_parts = []
            for idx, item in enumerate(analysis_result_list):
                ref_analysis_parts.append(f"Analysis ID {idx}: {item.brief_str()}")
            reference_analysis = "\n".join(ref_analysis_parts) if ref_analysis_parts else "No analysis results available"

            # Check for available chart files in output directory
            import os as _os
            target_name = self.config.config.get('target_name', 'Unknown')
            output_dir = _os.path.join(self.config.config.get('output_dir', './outputs'), target_name)
            ref_files = []
            if _os.path.isdir(output_dir):
                for root, dirs, files in _os.walk(output_dir):
                    for f in files:
                        if f.endswith(('.png', '.jpg', '.jpeg', '.svg', '.gif')):
                            rel_path = _os.path.relpath(_os.path.join(root, f), output_dir)
                            ref_files.append(f"  - {rel_path}")
            reference_image = "\n".join(ref_files) if ref_files else "No chart files available"

            return [{
                "role": "user",
                "content": self.SECTION_WRITING_PROMPT.format(
                    task=task,
                    report_theme=input_data.get('task'),
                    section_description=section_outline,
                    data_api=data_api_description,
                    data_info=data_info,
                    max_iterations=max_iterations,
                    target_language=self.target_language_name,
                    target_name=self.config.config.get('target_name', 'Unknown'),
                    stock_code=self.config.config.get('stock_code', 'Unknown'),
                    reference_data=reference_data,
                    reference_analysis=reference_analysis,
                    reference_image=reference_image
                )
            }]
        else:
            # Build reference strings for prompt placeholders
            ref_data_parts = []
            for idx, item in enumerate(collect_data_list):
                ref_data_parts.append(f"Data ID {idx}: {item.brief_str()}")
            reference_data = "\n".join(ref_data_parts) if ref_data_parts else "No collected data available"

            ref_analysis_parts = []
            for idx, item in enumerate(analysis_result_list):
                ref_analysis_parts.append(f"Analysis ID {idx}: {item.brief_str()}")
            reference_analysis = "\n".join(ref_analysis_parts) if ref_analysis_parts else "No analysis results available"

            return [{
                "role": "user",
                "content": self.SECTION_WRITING_WO_CHART_PROMPT.format(
                    task=task,
                    report_theme=input_data.get('task'),
                    section_description=section_outline,
                    data_api=data_api_description,
                    data_info=data_info,
                    max_iterations=max_iterations,
                    target_language=self.target_language_name,
                    target_name=self.config.config.get('target_name', 'Unknown'),
                    stock_code=self.config.config.get('stock_code', 'Unknown'),
                    reference_data=reference_data,
                    reference_analysis=reference_analysis
                )
            }]

    async def _handle_search_action(self, action_content: str):
        search_result = await self.tools[0].async_run(input_data={'query': action_content})
        return {
            'action': 'search',
            'action_content': action_content,
            'result': search_result['final_result'],
            'continue': True,
        }
    
    async def _handle_report_action(self, action_content: str):
        """Handle a 'final/report' action."""
        return {
            "action": "report",
            "action_content": action_content,
            "result": action_content,
            "continue": False,
        }
    async def _handle_outline_action(self, action_content: str):
        """Handle a 'outline' action."""
        return {
            "action": "outline",
            "action_content": action_content,
            "result": action_content,
            "continue": False,
        }
    
    async def _handle_draft_action(self, action_content: str):
        """Handle a 'outline' action."""
        return {
            "action": "draft",
            "action_content": action_content,
            "result": action_content,
            "continue": False,
        }
    
    async def _final_polish(self, section_input_data, draft_section: str):
        all_analysis_result = self.memory.get_analysis_result()
        all_image_list = []
        for analysis_result in all_analysis_result:
            all_image_list.extend(analysis_result.get_all_img())
        reference_image = '\n'.join(all_image_list)

        # Build reference_data and reference_analysis for prompt
        collect_data_list = self.memory.get_collect_data(exclude_type=['search', 'click'])
        ref_data_parts = []
        for idx, item in enumerate(collect_data_list):
            ref_data_parts.append(f"Data ID {idx}: {item.brief_str()}")
        reference_data = "\n".join(ref_data_parts) if ref_data_parts else "No collected data available"

        ref_analysis_parts = []
        for idx, item in enumerate(all_analysis_result):
            ref_analysis_parts.append(f"Analysis ID {idx}: {item.brief_str()}")
        reference_analysis = "\n".join(ref_analysis_parts) if ref_analysis_parts else "No analysis results available"

        final_prompt = self.FINAL_POLISH_PROMPT.format(
            draft_report = draft_section,
            reference_image = reference_image,
            reference_data = reference_data,
            reference_analysis = reference_analysis,
            target_language = self.target_language_name
        )
        
        final_message = [{"role": "user", "content": final_prompt}]
        output = await self.llm.generate(messages = final_message)
        final_section = extract_markdown(output)
        return final_section
    
    async def _replace_image_path(self, report):
        """
        Replace placeholder image references in the report with actual local paths.
        """
        # If charts are disabled, simply remove @import placeholders
        if not self.enable_chart:
            for section in report.sections:
                section_new_content = []
                for p_paragraph in section._content:
                    # Replace @import.* with empty string
                    p_paragraph = re.sub(r'@import.*', '', p_paragraph, flags=re.DOTALL)
                    section_new_content.append(p_paragraph)
                section._content = section_new_content
            return report
        
        def remove_suffix(name: str):
            return name.replace(".png", "").replace(".jpg", "").replace(".jpeg", "").replace(".md", "")
        def is_image_file(name: str):
            return name.endswith(".png") or name.endswith(".jpg") or name.endswith(".jpeg") or name.endswith(".md")
        all_analysis_result = self.memory.get_analysis_result()
        img_captions = []
        img_paths = []
        for analysis_result in all_analysis_result:
            short2long = {}
            img_dicts = {} # caption: abs_path 
            chart_name_mapping = analysis_result.chart_name_mapping
            for long_name, short_name in chart_name_mapping.items():
                short2long[remove_suffix(short_name)] = remove_suffix(long_name)
            image_save_dir = analysis_result.image_save_dir
            for image_name in os.listdir(image_save_dir):
                if is_image_file(image_name):
                    img_path = os.path.join(image_save_dir, image_name)
                    img_name = remove_suffix(image_name)
                    long_image_name = short2long.get(img_name, "")
                    if long_image_name != "":
                        img_dicts[long_image_name] = img_path
            img_captions.extend(list(img_dicts.keys()))
            img_paths.extend(list(img_dicts.values()))
        if len(img_captions) == 0:
            self.logger.warning("No image captions found from analysis results, attempting filesystem fallback...")
            # ── Filesystem fallback: scan all analyzer working dirs for images ──
            agent_working_dir = os.path.join(self.config.config.get('output_dir', './outputs'), self.config.config.get('target_name', 'Unknown'), "agent_working")
            if os.path.isdir(agent_working_dir):
                for agent_dir in sorted(os.listdir(agent_working_dir)):
                    if not agent_dir.startswith("agent_data_analyzer"):
                        continue
                    img_dir = os.path.join(agent_working_dir, agent_dir, "images")
                    if not os.path.isdir(img_dir):
                        continue
                    for img_filename in sorted(os.listdir(img_dir)):
                        if not is_image_file(img_filename):
                            continue
                        img_path = os.path.join(img_dir, img_filename)
                        caption = remove_suffix(img_filename)
                        img_captions.append(caption)
                        img_paths.append(img_path)
                        self.logger.info(f"  [fallback] Found: {img_filename}")

        if len(img_captions) == 0:
            self.logger.warning("No images found, replacing @import placeholders with fallback text")
            for section in report.sections:
                section_new_content = []
                for p_paragraph in section._content:
                    p_paragraph = re.sub(
                        r'@import.*',
                        '*[Chart not available — no chart images were generated during analysis]*',
                        p_paragraph,
                        flags=re.DOTALL,
                    )
                    section_new_content.append(p_paragraph)
                section._content = section_new_content
            return report
        self.logger.info(f"Building index for {len(img_captions)} images")

        # ── Phase 1: strip all @import placeholders ──
        for section in report.sections:
            cleaned = []
            for p in section._content:
                cleaned.append(re.sub(r'@import.*', '', p, flags=re.DOTALL))
            section._content = cleaned

        if len(img_captions) == 0:
            return report

        # ── Phase 2: build paragraph index & match each chart to its best paragraph ──
        para_entries = []  # [(section_idx, para_idx, text)]
        para_texts = []
        for si, section in enumerate(report.sections):
            for pi, p_text in enumerate(section._content):
                s = p_text.strip()
                if len(s) > 20:
                    para_entries.append((si, pi, s))
                    para_texts.append(s)

        if not para_texts:
            self.logger.warning("No paragraphs for chart placement — appending at end")
            for section in report.sections:
                for i, (cap, path) in enumerate(zip(img_captions, img_paths)):
                    section._content.append(get_md_img(path, cap, i + 1))
            return report

        try:
            para_index = IndexBuilder(config=self.config, embedding_model=self.use_embedding_name,
                                      working_dir=self.working_dir)
            await para_index._build_index(para_texts)
        except Exception as e:
            self.logger.warning(f"Para index failed: {e}; sequential fallback")
            ci = 0
            for section in report.sections:
                if ci < len(img_captions):
                    section._content.append(get_md_img(img_paths[ci], img_captions[ci], ci + 1))
                    ci += 1
            return report

        # insertion_plan[section_idx] = [(para_idx, img_markdown), ...]
        insertion_plan = {si: [] for si in range(len(report.sections))}
        used = set()

        for ci, caption in enumerate(img_captions):
            try:
                results = await para_index.search(caption)
                if not results:
                    continue
                best = int(results[0]['id'])
                if 0 <= best < len(para_entries):
                    si, pi, _ = para_entries[best]
                    insertion_plan[si].append((pi, get_md_img(img_paths[ci], caption, ci + 1)))
                    used.add(caption)
            except Exception as e:
                self.logger.warning(f"Chart '{caption[:40]}' placement failed: {e}")

        # ── Phase 3: rebuild sections with charts after their best paragraph ──
        for si, section in enumerate(report.sections):
            plan = sorted(insertion_plan.get(si, []), key=lambda x: x[0], reverse=True)
            for pi, img_md in plan:
                if pi + 1 < len(section._content):
                    section._content.insert(pi + 1, img_md)
                else:
                    section._content.append(img_md)

        # ── Phase 4: unmatched charts → end of last section ──
        unmatched = [(c, p) for c, p in zip(img_captions, img_paths) if c not in used]
        if unmatched and report.sections:
            base = len(used)
            for i, (cap, path) in enumerate(unmatched):
                report.sections[-1]._content.append(get_md_img(path, cap, base + i + 1))

        return report

    
    async def _add_abstract(self, input_data, report):
        """
        Add an abstract and update the title.
        """
        abstract_prompt = self.ABSTRACT_PROMPT
        title_prompt = self.TITLE_PROMPT


        response_content = await self.llm.generate(
            messages = [
            {
                'role': 'user',
                'content': abstract_prompt.format(target_language=self.target_language_name, report_content=report.content)
            }
        ])
        response_content = extract_markdown(response_content)
        report.abstract = response_content
        
        new_title = await self.llm.generate(
            messages = [
            {
                'role': 'user',
                'content': title_prompt.format(target_language=self.target_language_name, report_content=report.content)
            }
        ])
        new_title = new_title.replace("#","").strip()
        report._content = f"# {new_title}\n\n"

        return report

    async def _add_cover_page(self, input_data, report):
        pipeline_type = input_data.get('target_type', 'company')
        if pipeline_type != 'company':
            return report
        stock_code = input_data.get('stock_code', '')
        if stock_code == "":
            return report

        output_str = "\n\n## Company Fundamentals\n\n"
        # Three statements + shareholder profile
        collect_data_list = self.memory.get_collect_data()
        table_configs = [
            ("Income statement", "Income Statement"),
            ("Balance sheet", "Balance Sheet"),
            ("Cash-flow statement", "Cash-Flow Statement"),
            ("Shareholding structure", "Shareholder Structure"),
        ]
        for keyword, display_name in table_configs:
            target_item_list = [item for item in collect_data_list if keyword in item.name and stock_code in item.name]
            if len(target_item_list) == 0:
                print(f"No {display_name} data found")
                continue
            else:
                table_data = target_item_list[0].data
                if table_data is None:
                    print(f"{display_name} data is empty, skip formatting")
                    continue
                    
                if keyword in ["Income statement", "Balance sheet", "Cash-flow statement"]:
                    if 'Category' in table_data.columns:
                        table_data.rename(columns={'Category': 'Line item (RMB mn)'}, inplace=True)
                prompt = self.TABLE_BEAUTIFY_PROMPT.format(table_name=display_name, table_data=table_data.to_markdown(index=False))
                response = await self.llm.generate(
                    messages = [
                        {"role": "user", "content": prompt}
                    ]
                )
                table_string = "\n".join([line for line in response.split("\n") if line.strip() != ""])

                output_str += f'\n\n### {display_name}\n\n'
                output_str += table_string
                
                output_str += '\n\n'
        
        # Render stock-price chart
        try:
            self.logger.info("Rendering stock-price chart for cover page")
            target_item_list = [item for item in collect_data_list if 'candlestick' in item.name.lower() and stock_code in item.name]
            if len(target_item_list) != 0:
                kline_data = target_item_list[0].data
                if kline_data is None:
                    self.logger.warning("Candlestick data is empty; skip price visualization")
                else:
                    if isinstance(kline_data, list) and len(kline_data) == 1:
                        kline_data = kline_data[0]
                    if 'date' not in kline_data.columns:
                        if '\u65e5\u671f' in kline_data.columns:
                            kline_data.rename(columns={'\u65e5\u671f': 'date'}, inplace=True)
                        if '\u6536\u76d8' in kline_data.columns:
                            kline_data.rename(columns={'\u6536\u76d8': 'close'}, inplace=True)
                    fig_path = draw_kline_chart(kline_data, self.working_dir)
                    output_str += f'\n\n### Share Price Trend\n\n'
                    output_str += f'![Trailing price performance]({fig_path})\n\n'
        except Exception as e:
            self.logger.error(f"Failed to draw price trend: {e}", exc_info=True)
            pass

        first_section = Section('Company Fundamentals', output_str)
        first_section.set_content(output_str)
        report.sections = [first_section] + report.sections

        return report
    

    async def _add_reference(self, report):
        """
        Append the reference-data section and replace placeholder citations.
        """
        collect_data_list = self.memory.get_collect_data() # only use data, without analysis result
        all_data = []
        for item in collect_data_list:
            # Keep compatibility with ToolResult fields expected by downstream report assembly.
            name = item.name + '\n' + item.description # used for index
            content = item.source # used for display citation
            # for url, find the title in search results
            if isinstance(item, ClickResult):
                url = item.link
                title = self.memory.get_url_title(url)
                if title == "":
                    title = item.name
                content = f"{title}\n{url}"

            # content = item.name + '\n' + item.link  # used for display citation
            if content not in [ii['content'] for ii in all_data]:
                all_data.append({
                    'name': name,
                    'content': content 
                })
        self.logger.info(f"Total data for reference: {len(all_data)}")

        total_corpus = [item['name'] for item in all_data]
        index = IndexBuilder(config=self.config, embedding_model=self.use_embedding_name, working_dir=self.working_dir)
        try:
            await index._build_index(total_corpus)
        except Exception as e:
            self.logger.warning(f"Failed to build embedding index: {e}. References will use fallback numbering.")
            # Fallback: assign sequential numbers to all references without semantic matching
            reference_str = "## Reference Data Sources\n\n"
            for idx, item in enumerate(all_data):
                content = item['content'].replace("[PDF]", "")
                parts = content.strip().split("\n")
                if len(parts) >= 2 and parts[-1].strip().startswith("http"):
                    title = " ".join(parts[:-1]).strip()
                    url = parts[-1].strip()
                    reference_str += f"{idx + 1}. [{title}]({url})\n"
                elif " http" in content:
                    i = content.index(" http")
                    reference_str += f"{idx + 1}. [{content[:i].strip()}]({content[i:].strip()})\n"
                else:
                    reference_str += f"{idx + 1}. {content}\n"
            new_section = Section('Reference Data Sources', reference_str)
            new_section.set_content(reference_str)
            report.sections.append(new_section)
            return report

        total_cited_dict = {}
        for section in report.sections:
            # Optional: log section length
            try:
                self.logger.debug(f"Processing section, content length={len(section.content)}")
            except Exception:
                pass
            section_new_content = []
            for p_paragraph in section._content:
                content = p_paragraph
                # Locate citation placeholders
                match_list = re.findall(r'\[[Ss]ource[：:]\s*(.*?)\]',content)
                self.logger.debug(f"Match list: {match_list}")
                for match_item in match_list:
                    # Use BM25/embedding search
                    search_result = await index.search(match_item, top_k=5)
                    if not search_result:
                        continue  # Skip this citation if search returns empty
                    score_list = [item['score'] for item in search_result]
                    id_list = [item['id'] for item in search_result]  # Get actual data indices
                    self.logger.debug(f"Score list: {score_list}")
                    self.logger.debug(f"ID list: {id_list}")
                    # Sort by score (descending) and get corresponding indices
                    sorted_idx = np.argsort(score_list)[::-1]
                    score_list = np.array(score_list)
                    score_list = np.exp(score_list) / np.sum(np.exp(score_list))

                    cite_list = []
                    for pos in sorted_idx:
                        pos = int(pos)
                        actual_idx = id_list[pos]  # Get the actual data index
                        if score_list[pos] > 0.2 and len(cite_list) < 5:
                            cite_list.append(actual_idx)
                    if len(cite_list) == 0:
                        # If no item meets threshold, use the top result
                        cite_list.append(id_list[sorted_idx[0]])
                    new_cite_list = []
                    for idx in cite_list:
                        if idx not in total_cited_dict:
                            total_cited_dict[idx] = len(total_cited_dict) + 1
                    new_cite_list = [total_cited_dict[idx] for idx in cite_list]
                    # Build the regex for replacement
                    pattern_to_replace = r'\[[Ss]ource[：:]\s*' + re.escape(match_item) + r'\]'
                    content = re.sub(pattern_to_replace, f'[{",".join([str(item) for item in new_cite_list])}]', content)

                section_new_content.append(content)
            section._content = section_new_content


        reference_str = "## Reference Data Sources\n\n"
        for old_index, new_index in total_cited_dict.items():
            content = all_data[old_index]['content']
            content = content.replace("[PDF]", "")
            # Parse "Title\nURL" or "Title URL" into clickable markdown link
            parts = content.strip().split("\n")
            if len(parts) >= 2 and parts[-1].strip().startswith("http"):
                title = " ".join(parts[:-1]).strip()
                url = parts[-1].strip()
                reference_str += f"{new_index}. [{title}]({url})\n"
            elif " http" in content:
                # "Title http://..." format
                idx = content.index(" http")
                title = content[:idx].strip()
                url = content[idx:].strip()
                reference_str += f"{new_index}. [{title}]({url})\n"
            else:
                reference_str += f"{new_index}. {content}\n"
        new_section = Section('Reference Data Sources', reference_str)
        new_section.set_content(reference_str)
        report.sections.append(new_section)
        return report

        

    async def post_process_report(self, input_data, report):
        """
        Post-process the report while saving progress between sub-stages:
          0: replace image paths
          1: add abstract and title
          2: add cover/basic data page
          3: add reference data section
          4: render to docx
        """
        current_state = {
            'phase': 'post_process',
            'post_stage': self._post_stage,
            'report_obj': report,
        }
        # 0 Replace image paths
        if self._post_stage <= 0:
            self.logger.info("[Phase2] Step 0: replace image paths")
            report = await self._replace_image_path(report)
            self._post_stage = 1
            current_state['report_obj_stage1'] = copy.deepcopy(report)
            current_state['report_obj'] = report
            current_state['post_stage'] = self._post_stage
            await self.save(state=current_state, checkpoint_name='report_latest.pkl')
            self.logger.info("[Phase2] Step 0 done, checkpoint saved")

        # 1 Add abstract/title (conditional based on add_introduction setting)
        if self._post_stage <= 1:
            if getattr(self, 'add_introduction', True):
                self.logger.info("[Phase2] Step 1: add abstract and title")
                report = await self._add_abstract(input_data, report)
            else:
                self.logger.info("[Phase2] Step 1: skipping abstract/introduction (add_introduction=False for general reports)")
                # Still generate a better title
                new_title = await self.llm.generate(
                    messages = [
                    {
                        'role': 'user',
                        'content': self.TITLE_PROMPT.format(target_language=self.target_language_name, report_content=report.content)
                    }
                ])
                new_title = new_title.replace("#","").strip()
                report._content = f"# {new_title}\n\n"
            self._post_stage = 2
            current_state['report_obj_stage2'] = copy.deepcopy(report)
            current_state['report_obj'] = report
            current_state['post_stage'] = self._post_stage
            await self.save(state=current_state, checkpoint_name='report_latest.pkl')
            self.logger.info("[Phase2] Step 1 done, checkpoint saved")

        # 2 Add cover/basic data page
        if self._post_stage <= 2:
            self.logger.info("[Phase2] Step 2: add cover/basic data page")
            report = await self._add_cover_page(input_data, report)
            self._post_stage = 3
            current_state['report_obj_stage3'] = copy.deepcopy(report)
            current_state['report_obj'] = report
            current_state['post_stage'] = self._post_stage
            await self.save(state=current_state, checkpoint_name='report_latest.pkl')
            self.logger.info("[Phase2] Step 2 done, checkpoint saved")

        # 3 Add references (conditional based on add_reference_section setting)
        if self._post_stage <= 3:
            if getattr(self, 'add_reference_section', True):
                self.logger.info("[Phase2] Step 3: add references")
                report = await self._add_reference(report)
            else:
                self.logger.info("[Phase2] Step 3: skipping reference section (add_reference_section=False)")
            self._post_stage = 4
            current_state['report_obj_stage4'] = copy.deepcopy(report)
            current_state['report_obj'] = report
            current_state['post_stage'] = self._post_stage
            await self.save(state=current_state, checkpoint_name='report_latest.pkl')
            self.logger.info("[Phase2] Step 3 done, checkpoint saved")

        # 4 Render to docx
        if self._post_stage <= 4:
            self.logger.info("[Phase2] Step 4: render report to docx")
            working_dir = self.config.config['working_dir']
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            md_path = os.path.join(working_dir, f'{report.title}_{timestamp}.md')
            docx_path = os.path.join(working_dir, f'{report.title}_{timestamp}.docx')
            content = report.content
            content = content.replace("```markdown", "").replace("```", "")
            media_dir = os.path.join(working_dir, "media")

            # ── Copy referenced images to media/ with safe filenames, rewrite paths to relative ──
            import shutil
            img_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
            md_dir = os.path.dirname(md_path)
            os.makedirs(media_dir, exist_ok=True)
            copied_count = 0
            def replace_img_ref(match):
                nonlocal copied_count
                alt_text = match.group(1)
                src_path = match.group(2).strip()
                if not os.path.isabs(src_path) or not os.path.exists(src_path):
                    return match.group(0)
                # Sanitize filename: keep only safe chars, use .png extension
                safe_name = f"chart_{copied_count + 1:02d}.png"
                dest_path = os.path.join(media_dir, safe_name)
                try:
                    if not os.path.exists(dest_path):
                        shutil.copy2(src_path, dest_path)
                    copied_count += 1
                    self.logger.info(f"  Copied image to media/: {safe_name} <- {os.path.basename(src_path)}")
                except Exception as e:
                    self.logger.warning(f"  Failed to copy image {src_path}: {e}")
                    return match.group(0)
                return f"![{alt_text}](media/{safe_name})"
            content = img_pattern.sub(replace_img_ref, content)
            self.logger.info(f"Copied {copied_count} images to media/ for DOCX embedding")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(content)
            reference_doc = self.config.config.get('reference_doc_path', None)
            # Auto-detect template from skill's bundled template directory
            if not reference_doc or not os.path.exists(reference_doc):
                template_dir = Path(__file__).resolve().parent.parent.parent / "template"
                default_template = template_dir / "report_template_backup.docx"
                if default_template.exists():
                    reference_doc = str(default_template)
                    self.logger.info(f"Using default template: {reference_doc}")
                else:
                    self.logger.warning(f"Default template not found at {default_template}")

            # Use pypandoc (bundled pandoc binary) for md → docx conversion
            import pypandoc
            extra_args = [
                "--standalone",
                "--toc",
                "--toc-depth=3",
                f"--resource-path={working_dir}",
            ]
            if reference_doc and os.path.exists(reference_doc):
                extra_args.append(f"--reference-doc={reference_doc}")
            if os.path.exists(media_dir):
                extra_args.append(f"--extract-media={media_dir}")

            try:
                pypandoc.convert_file(
                    md_path, 'docx',
                    outputfile=docx_path,
                    extra_args=extra_args,
                )
                self.logger.info(f"DOCX generated: {docx_path}")

                # Inject real Word TOC field (pandoc --toc is static; this adds updatable page numbers)
                try:
                    _inject_word_toc(docx_path)
                    self.logger.info(f"TOC field injected: {docx_path}")
                except Exception as e:
                    self.logger.warning(f"TOC injection failed (non-fatal): {e}")

                # Normalize every Word text part and fallback to the central Latin font.
                try:
                    font_stats = normalize_docx_typography(docx_path)
                    self.logger.info(
                        f"DOCX typography normalized to {LATIN_FONT}: "
                        f"{font_stats} ({docx_path})"
                    )
                except Exception as e:
                    if os.path.exists(docx_path):
                        os.remove(docx_path)
                    raise RuntimeError("DOCX typography normalization failed") from e
            except Exception as e:
                self.logger.error(f"Failed to convert md to docx: {e}", exc_info=True)
            
            # Validate output is non-empty
            if not os.path.exists(md_path) or os.path.getsize(md_path) == 0:
                self.logger.error(
                    f"Report output is empty: {md_path}. "
                    "This usually means sections produced no content. "
                    "Check section generation logs for errors."
                )

            pdf_path = docx_path.replace(".docx", ".pdf")
            pdf_generated = False

            # Strategy 1: docx2pdf (requires Microsoft Word on Windows)
            try:
                if _HAS_DOCX2PDF:
                    docx2pdf.convert(docx_path, pdf_path)
                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                        pdf_generated = True
                        self.logger.info(f"PDF generated via docx2pdf: {pdf_path}")
                else:
                    self.logger.info("docx2pdf not available; trying pandoc for PDF...")
            except Exception as e:
                self.logger.warning(f"docx2pdf failed: {e}; trying pandoc fallback...")

            # Strategy 2: pandoc MD → PDF direct (requires pdflatex or wkhtmltopdf)
            if not pdf_generated:
                try:
                    import pypandoc
                    pdf_extra_args = [
                        "--standalone",
                        "--toc",
                        "--toc-depth=3",
                        f"--resource-path={working_dir}",
                        "--pdf-engine=xelatex",
                        f"--variable=mainfont:{LATIN_FONT}",
                        f"--variable=sansfont:{LATIN_FONT}",
                        f"--variable=monofont:{LATIN_FONT}",
                    ]
                    pypandoc.convert_file(
                        md_path, 'pdf',
                        outputfile=pdf_path,
                        extra_args=pdf_extra_args,
                    )
                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                        pdf_generated = True
                        self.logger.info(f"PDF generated via pandoc: {pdf_path}")
                except Exception as e:
                    self.logger.warning(f"Pandoc PDF also failed: {e}")

            if not pdf_generated:
                self.logger.warning(
                    "PDF not generated. Install one of: Microsoft Word (for docx2pdf), "
                    "or MiKTeX/TeX Live (for pandoc xelatex), or wkhtmltopdf."
                )
            self._post_stage = 5
            current_state['rendered_md'] = md_path
            current_state['rendered_docx'] = docx_path
            current_state['finished'] = True
            await self.save(state=current_state, checkpoint_name='report_latest.pkl')
            self.logger.info(f"[Phase2] Step 4 done, rendered files: md={md_path}, docx={docx_path}, pdf={pdf_path}")
        return report

    def _get_persist_extra_state(self) -> Dict[str, Any]:
        """
        Provide extra state so parent persistence can restore the pipeline stages.
        """
        return {
            'phase': getattr(self, '_phase', 'outline'),
            'section_index': getattr(self, '_section_index_done', 0),
            'post_stage': getattr(self, '_post_stage', 0),
        }

    def _load_persist_extra_state(self, state: Dict[str, Any]):
        """
        Restore stage metadata from a checkpoint.
        """
        # Recover from the extra field populated by _get_persist_extra_state
        extra = state.get('extra', {})
        
        phase = extra.get('phase') or state.get('phase')
        if isinstance(phase, str):
            self._phase = phase
        
        section_index = extra.get('section_index') or state.get('section_index')
        if section_index is not None:
            try:
                self._section_index_done = int(section_index)
            except Exception:
                pass
        
        post_stage = extra.get('post_stage') or state.get('post_stage')
        if post_stage is not None:
            try:
                self._post_stage = int(post_stage)
            except Exception:
                pass
        
        enable_chart = extra.get('enable_chart') or state.get('enable_chart')
        if enable_chart is not None:
            try:
                self.enable_chart = bool(enable_chart)
            except Exception:
                pass
        else:
            self.enable_chart = True
    
    async def _prepare_outline_prompt(self, input_data):
        max_iterations = input_data.get('max_iterations', 10)
        outline_template_path = self.config.config.get('outline_template_path', None)
        
        if outline_template_path is None or not os.path.exists(outline_template_path):
            outline_template = ""
        else:
            with open(outline_template_path, 'r', encoding='utf-8') as f:
                outline_template = f.read()
        # Prepare data API description and available analysis info
        data_api_description = self.prompt_loader.get_prompt('data_api_outline')
        analysis_result_list = self.memory.get_analysis_result()
        
        data_info = "You have access to the following analysis results:\n\n"
        for idx, result in enumerate(analysis_result_list):
            data_info += f"**Analysis Report ID {idx}:**\n{result.brief_str()}\n\n"
        data_info += "\nYou can retrieve detailed content using `get_analysis_result(analysis_id)` in your code.\n"
        
        initial_prompt = self.DRAFT_GENERATOR_PROMPT.format(
            task=input_data['task'],
            report_requirements=outline_template,
            data_api=data_api_description,
            data_info=data_info,
            max_iterations=max_iterations,
            target_language=self.target_language_name,
            target_name=self.config.config.get('target_name', 'Unknown'),
            stock_code=self.config.config.get('stock_code', 'Unknown')
        )
        return [{"role": "user", "content": initial_prompt}]

    async def generate_outline(
        self, 
        input_data, 
        max_iterations: int = 10,
        stop_words: list[str] = [],
        echo=False,
        resume: bool = True,
        checkpoint_name: str = 'outline_latest.pkl'
    ):
        """
        Generate the report outline via agentic workflow.

        Args:
            input_data: Dict containing task metadata.
            max_iterations: Maximum number of interaction rounds.

        Returns:
            Report object populated with outline sections.
        """
       
        # Prepare executor for outline generation
        await self._prepare_executor()

        self.logger.info(f"[Outline] Starting agentic outline generation (max {max_iterations} rounds)")
        
        # Create input data for outline generation
        outline_input_data = {
            'task': input_data['task'],
            'max_iterations': max_iterations
        }
        self.current_task_data = outline_input_data

        outline_result = await super().async_run(
            input_data=outline_input_data,
            max_iterations=max_iterations,
            stop_words=stop_words,
            echo=echo,
            resume=resume,
            checkpoint_name=checkpoint_name,
            prompt_function=self._prepare_outline_prompt,
        )
    
        outline_content = extract_markdown(outline_result['final_result'])
        
        return Report(outline_content) if outline_content else Report("# Error: Could not generate outline")



    async def async_run(
        self, 
        input_data: dict, 
        max_iterations: int = 10,
        stop_words: list[str] = [],
        # stop_words: list[str] = ["</draft>", "</outline>", "</report>", "</execute>"],
        echo=False,
        resume: bool = True,
        checkpoint_name: str = 'report_latest.pkl',
        enable_chart = True,
        add_introduction: bool = None,  # None means auto-detect based on target_type
        add_reference_section: bool = True
    ) -> dict:
        """
        Three-stage execution flow for the report generator:
        Phase 0: outline creation
        Phase 1: per-section drafting
        Phase 2: post processing
        """
        # Initialize/restore stage state
        report = None
        start_index = 0
        self.enable_chart = enable_chart
        input_data['max_iterations'] = max_iterations
        
        # Configure post-processing options based on target_type
        target_type = self.config.config.get('target_type', 'general')
        
        # For general/deep-research reports, default to NO introduction (user specifies their own structure)
        # For company/financial reports, default to YES (standard report format)
        if add_introduction is None:
            self.add_introduction = target_type not in ['general']
        else:
            self.add_introduction = add_introduction
        
        self.add_reference_section = add_reference_section
        
        if resume:
            state = await self.load(checkpoint_name=checkpoint_name)
            if state is not None:
                # Restore extra metadata
                self._load_persist_extra_state(state)
                self.logger.info(f"[Resume] phase={getattr(self, '_phase', None)}, section_index={getattr(self, '_section_index_done', None)}, post_stage={getattr(self, '_post_stage', None)}")
                
                # If the workflow already finished, return the saved report
                if state.get('finished'):
                    restored_report = state.get('report_obj')
                    if restored_report:
                        self.logger.info("Report already completed, restoring from checkpoint")
                        return restored_report
                
                # Restore an in-progress report if available
                restored_report = state.get('report_obj')
                if restored_report is not None:
                    report = restored_report
                    start_index = self._section_index_done
                    self.logger.info(f"[Resume] Restored report object, will resume from section_index={start_index}")
        
        # Phase 0: outline generation
        if self._phase == 'outline' or report is None:
            self.logger.info("[Phase0] Generating Report Outline")
            report = await self.generate_outline(
                input_data, 
                max_iterations=max_iterations,
                stop_words=stop_words,
                echo=echo,
                resume=resume,
                checkpoint_name='outline_latest.pkl'
            )
            self._phase = 'sections'
            # Persist outline state
            await self.save(
                state={
                    'phase': self._phase,
                    'report_obj': report,
                    'input_data': input_data,
                    'enable_chart': self.enable_chart,
                },
                checkpoint_name=checkpoint_name,
            )
            self.memory.save()
            self.logger.info(f"[Phase0] Completed: outline sections={len(report.sections)}")

        
        # Phase 1: per-section generation
        if self._phase == 'sections':
            self.logger.info("[Phase1] Begin generating sections")
            # Section generation is sequential to preserve report order and cache behavior.
            for idx, section in enumerate(report.sections):
                if idx < start_index:
                    continue
                section_input_data = input_data.copy()
                section_input_data['section_outline'] = section.outline
                self.logger.info(f"[Phase1] Section {idx+1}/{len(report.sections)} start")
                
                # Prepare executor with data access functions for agentic workflow
                await self._prepare_executor()
                
                # Each section run has its own checkpoint for resume support
                section_result = await super().async_run(
                    input_data=section_input_data,
                    max_iterations=max_iterations,
                    stop_words=stop_words,
                    echo=echo,
                    resume=resume and idx == start_index,
                    checkpoint_name=f'section_{idx}.pkl'
                )
                draft_section = section_result['final_result']
                self.logger.debug(f"[Phase1] Draft section length={len(draft_section)}")
                
                # Final polish for the section content
                final_section = await self._final_polish(section_input_data, draft_section)
                self.logger.debug(f"[Phase1] Final section length={len(final_section)}")
                self.memory.add_log(
                    id=self.id,
                    type=self.type,
                    input_data=section_input_data,
                    output_data=section_result,
                    error=False,
                    note=f"Report generator executed successfully"
                )
                section.set_content(final_section) 
                # Save global progress after each section to resume later
                await self.save(
                    state={
                        'phase': 'sections',
                        'section_index': idx + 1,
                        'report_obj': report,
                        'input_data': input_data,
                    },
                    checkpoint_name=checkpoint_name,
                )
                self.memory.save()
                # Update in-memory progress pointer
                self._section_index_done = idx + 1
                self.logger.info(f"[Phase1] Section {idx+1} done, checkpoint saved (section_index={self._section_index_done})")
            
            # Move to post-process stage once all sections are done
            self._phase = 'post_process'
            await self.save(
                state={
                    'phase': self._phase,
                    'section_index': self._section_index_done,
                    'post_stage': self._post_stage,
                    'report_obj': report,
                    'input_data': input_data,
                },
                checkpoint_name=checkpoint_name,
            )
            self.memory.save()
            self.logger.info("[Phase1] Completed: All sections generated")

        # Phase 2: post processing (resumable)
        if self._phase == 'post_process':
            self.logger.info("[Phase2] Begin post processing")
            report = await self.post_process_report(input_data, report)
            self.memory.save()
            self.logger.info("[Phase2] Completed post processing")

        return report
    
