"""Regression tests for centralized report typography."""

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


SCRIPTS_DIR = Path(__file__).resolve().parents[1]
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from src.typography import (  # noqa: E402
    LATIN_FONT,
    apply_typography_tokens,
    configure_matplotlib_fonts,
    enforce_figure_text_fonts,
    matplotlib_setup_code,
    normalize_docx_typography,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS = {"w": W_NS, "a": A_NS}


class TypographyTests(unittest.TestCase):
    def _build_fixture(self, path: Path) -> None:
        doc = Document()
        heading = doc.add_heading("Revenue 2026 增长", level=1)
        heading.runs[0].font.name = "Arial"

        paragraph = doc.add_paragraph()
        run = paragraph.add_run("EBITDA 增长 12.5%")
        r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_fonts.set(qn("w:asciiTheme"), "majorHAnsi")
        r_fonts.set(qn("w:hAnsiTheme"), "minorHAnsi")
        r_fonts.set(qn("w:hint"), "eastAsia")
        r_fonts.set(qn("w:eastAsia"), "楷体")

        doc.add_paragraph("Numbered 123", style="List Number")
        doc.sections[0].header.paragraphs[0].add_run("Header 2026")
        footer = doc.sections[0].footer.paragraphs[0]
        footer.add_run("Page ")
        field_run = OxmlElement("w:r")
        field_text = OxmlElement("w:t")
        field_text.text = "7"
        field_run.append(field_text)
        footer._element.append(field_run)
        doc.save(path)

    def _read_xml(self, archive: zipfile.ZipFile, filename: str):
        return etree.fromstring(archive.read(filename))

    def test_docx_normalization_covers_all_word_text_parts(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "font_fixture.docx"
            self._build_fixture(path)

            stats = normalize_docx_typography(path)
            self.assertGreater(stats["runs"], 0)
            self.assertGreater(stats["styles"], 0)
            self.assertGreater(stats["numbering_levels"], 0)

            with zipfile.ZipFile(path) as archive:
                for filename in archive.namelist():
                    if not filename.startswith("word/") or not filename.endswith(".xml"):
                        continue
                    root = self._read_xml(archive, filename)
                    for r_fonts in root.xpath(".//w:rFonts", namespaces=NS):
                        self.assertEqual(r_fonts.get(qn("w:ascii")), LATIN_FONT, filename)
                        self.assertEqual(r_fonts.get(qn("w:hAnsi")), LATIN_FONT, filename)
                        self.assertIsNone(r_fonts.get(qn("w:asciiTheme")), filename)
                        self.assertIsNone(r_fonts.get(qn("w:hAnsiTheme")), filename)
                        self.assertIsNone(r_fonts.get(qn("w:hint")), filename)

                    for run in root.xpath(".//w:r", namespaces=NS):
                        r_fonts = run.find("w:rPr/w:rFonts", namespaces=NS)
                        self.assertIsNotNone(r_fonts, filename)
                        self.assertEqual(r_fonts.get(qn("w:ascii")), LATIN_FONT, filename)

                styles = self._read_xml(archive, "word/styles.xml")
                for style in styles.findall("w:style", namespaces=NS):
                    r_fonts = style.find("w:rPr/w:rFonts", namespaces=NS)
                    self.assertIsNotNone(r_fonts)
                    self.assertEqual(r_fonts.get(qn("w:ascii")), LATIN_FONT)

                numbering = self._read_xml(archive, "word/numbering.xml")
                for level in numbering.xpath(".//w:lvl", namespaces=NS):
                    r_fonts = level.find("w:rPr/w:rFonts", namespaces=NS)
                    self.assertIsNotNone(r_fonts)
                    self.assertEqual(r_fonts.get(qn("w:ascii")), LATIN_FONT)

                theme = self._read_xml(archive, "word/theme/theme1.xml")
                for latin in theme.xpath(".//a:latin", namespaces=NS):
                    self.assertEqual(latin.get("typeface"), LATIN_FONT)

                document = self._read_xml(archive, "word/document.xml")
                mixed_run = document.xpath(".//w:r[w:t[contains(., 'EBITDA')]]", namespaces=NS)[0]
                mixed_fonts = mixed_run.find("w:rPr/w:rFonts", namespaces=NS)
                self.assertEqual(mixed_fonts.get(qn("w:eastAsia")), "楷体")

            reopened = Document(path)
            self.assertIn("EBITDA 增长 12.5%", "\n".join(p.text for p in reopened.paragraphs))

    def test_prompt_and_matplotlib_settings_share_the_central_font(self):
        rendered = apply_typography_tokens(
            "Latin=__LATIN_FONT__; families=__CHART_FONT_FAMILIES__"
        )
        self.assertIn(LATIN_FONT, rendered)
        self.assertNotIn("__LATIN_FONT__", rendered)

        module = SimpleNamespace(rcParams={})
        families = configure_matplotlib_fonts(module)
        self.assertEqual(families[0], LATIN_FONT)
        self.assertEqual(module.rcParams["font.family"][0], LATIN_FONT)
        setup_code = matplotlib_setup_code()
        self.assertIn(LATIN_FONT, setup_code)
        compile(setup_code, "<matplotlib-setup>", "exec")

        class Artist:
            def set_fontfamily(self, value):
                self.family = value

        artist = Artist()
        figure = SimpleNamespace(findobj=lambda: [artist, object()])
        self.assertEqual(enforce_figure_text_fonts(figure), 1)
        self.assertEqual(artist.family[0], LATIN_FONT)


if __name__ == "__main__":
    unittest.main()
